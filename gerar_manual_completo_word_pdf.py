import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

import os
import shutil

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak, HRFlowable, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

desktop_path = r"C:\Users\Otávio\Desktop"
os.makedirs(desktop_path, exist_ok=True)

docx_desktop = os.path.join(desktop_path, "Manual_do_Analista_Lincros.docx")
pdf_desktop = os.path.join(desktop_path, "Manual_do_Analista_Lincros.pdf")

docx_local = "Manual_do_Analista_Lincros.docx"
pdf_local = "Manual_do_Analista_Lincros.pdf"

# =============================================================================
# 1. GERAR DOCUMENTO WORD (.DOCX) COMPLETO E DETALHADO
# =============================================================================
doc = docx.Document()

# Margens do Documento (0.75 polegadas)
for s in doc.sections:
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.75)
    s.right_margin = Inches(0.75)

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def add_callout_box(doc, title, text, bg_color="FEF08A", border_title="📌 MARCA-TEXTO DE VERIFICAÇÃO"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_color)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    r_title = p.add_run(f"{border_title}: {title}\n")
    r_title.bold = True
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = RGBColor(133, 77, 14)
    
    r_text = p.add_run(text)
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(66, 32, 6)
    
    doc.add_paragraph()

# Cabecalho do Documento
p_t = doc.add_paragraph()
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = p_t.add_run("MANUAL DO ANALISTA DE LOGÍSTICA")
rt.bold = True
rt.font.size = Pt(24)
rt.font.color.rgb = RGBColor(26, 54, 93)

p_s = doc.add_paragraph()
p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = p_s.add_run("Guia Detalhado de Padronização de Planilhas e Operação no Modelo Lincros\n")
rs.italic = True
rs.font.size = Pt(13)
rs.font.color.rgb = RGBColor(74, 85, 104)

doc.add_paragraph()

# -----------------------------------------------------------------------------
# CAPÍTULO 1: FLUXO DE TRABALHO
# -----------------------------------------------------------------------------
doc.add_heading("1. Visão Geral e Fluxo Recomendado de Trabalho", level=1)
doc.add_paragraph(
    "Este manual descreve detalhadamente o processo de preparação, conversão e estrutura exigida para cadastrar tabelas de frete, faixas de CEP, prazos e rotas no sistema Lincros."
)

doc.add_paragraph(
    "Para evitar falhas de importação ou rejeição de arquivos na plataforma Lincros, os analistas devem seguir a sequência de 5 passos abaixo:"
)

passos = [
    ("Passo 1: Baixar a Planilha Modelo Base", "Acesse a barra lateral do aplicativo e clique no botão 'Baixar Modelo Base (Vazio)'. Este arquivo já contém todos os cabeçalhos reconhecidos pelas ferramentas do sistema."),
    ("Passo 2: Tratar Localidades (CEP ou IBGE)", "Se o contrato for por faixas de CEP, utilize a aba 'Cadastro CEP'. Se o contrato for por municípios, utilize a aba 'Preencher IBGE' para garantir o código oficial de 7 dígitos."),
    ("Passo 3: Definir Prazos e Frequências", "Preencha a coluna de prazos em dias úteis e marque a frequência dos dias da semana (Segunda a Sábado) com a indicação de atendimento."),
    ("Passo 4: Gerar o Modelo Região Lincros", "Suba a planilha base tratada na aba 'Criar Região'. O sistema construirá a estrutura oficial com as abas 'regioes' e 'localizacoes_atendidas'."),
    ("Passo 5: Estruturar a Matriz de Rotas", "Na aba 'Gerar Rotas', envie a planilha de regiões gerada no passo anterior e vincule a origem para criar a matriz final de rotas.")
]

for title, desc in passos:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(f"{title}: ")
    r1.bold = True
    p.add_run(desc)

add_callout_box(
    doc,
    "CONFIGURAÇÕES OBRIGATÓRIAS DA BARRA LATERAL",
    "Antes de clicar nos botões de processamento, verifique sempre se o CNPJ e o Nome da Transportadora estão devidamente preenchidos na barra lateral esquerda. O sistema utiliza esses valores para vincular as unidades atendidas nas tabelas Lincros.",
    "FEF08A"
)

# Inserir Figura 1
doc.add_heading("1.1 Estrutura da Planilha Modelo Base Preenchida pelo Analista", level=2)
doc.add_paragraph("Abaixo é ilustrado o modelo base completo (`Base_de_Origem_Template.xlsx`), destacando a presença das novas colunas de CEP Inicial e CEP Final ao lado do Código IBGE:")

if os.path.exists("manual_assets/fig01_base_analista.png"):
    doc.add_picture("manual_assets/fig01_base_analista.png", width=Inches(6.5))

# -----------------------------------------------------------------------------
# CAPÍTULO 2: CADASTRO CEP
# -----------------------------------------------------------------------------
doc.add_heading("2. Módulo: Cadastro CEP (Lincros)", level=1)
doc.add_paragraph(
    "O módulo Cadastro CEP permite ao analista colar uma lista de CEPs ou enviar uma planilha sem formatar. O sistema realiza consultas em tempo real nas APIs oficiais dos Correios (ViaCEP, AwesomeAPI, ApiCEP e BrasilAPI) utilizando multithreading paralelo para processar até 10.000 linhas em poucos segundos."
)

doc.add_heading("2.1 Como o Analista deve preparar os dados de entrada", level=2)
doc.add_paragraph(
    "O analista pode colar a lista diretamente nas caixas de texto ou subir uma planilha Excel/CSV. O sistema aceita CEPs com hífen (`80010-000`) ou sem hífen (`80010000`). Se a coluna CEP Final não for informada, o sistema replica automaticamente o CEP Inicial."
)

if os.path.exists("manual_assets/fig02_cep_entrada.png"):
    doc.add_picture("manual_assets/fig02_cep_entrada.png", width=Inches(6.2))

doc.add_heading("2.2 Como o Sistema retorna o Modelo CEP Lincros", level=2)
doc.add_paragraph(
    "Ao clicar em 'PROCESSAR CEPS & CONSULTAR CORREIOS', o sistema gera a planilha 'Modelo CEP Preenchido.xlsx' no padrão estrito exigido pelo Lincros:"
)

if os.path.exists("manual_assets/fig03_cep_retorno.png"):
    doc.add_picture("manual_assets/fig03_cep_retorno.png", width=Inches(6.5))

add_callout_box(
    doc,
    "DETALHAMENTO DAS COLUNAS GERADAS NO MODELO CEP",
    "• Coluna A (ID Localização): Mantida completamente vazia para ser gerada pela Lincros na importação.\n"
    "• Coluna B (CEP Inicial) e Coluna C (CEP Final): Formato numérico de 8 dígitos limpos (ex: 80010000).\n"
    "• Coluna D (Nome): Preenchida automaticamente no padrão 'Cidade - UF' (ex: Curitiba - PR).\n"
    "• Coluna E (Ativo): Preenchida exatamente como 'VERDADEIRO'.",
    "FEF08A"
)

# -----------------------------------------------------------------------------
# CAPÍTULO 3: PREENCHER IBGE
# -----------------------------------------------------------------------------
doc.add_heading("3. Módulo: Preencher Códigos IBGE", level=1)
doc.add_paragraph(
    "Esta ferramenta lê a planilha base enviada pelo analista e localiza os códigos IBGE oficiais de 7 dígitos para cada cidade, evitando que o analista tenha que procurar manualmente código por código."
)

doc.add_heading("3.1 Requisitos de Entrada e Retorno Gerado", level=2)
doc.add_paragraph(
    "A planilha do analista deve possuir as colunas chamadas 'Destino' (nome da cidade) e 'UF Destino' (sigla do estado com 2 letras)."
)

if os.path.exists("manual_assets/fig04_ibge_entrada.png"):
    doc.add_picture("manual_assets/fig04_ibge_entrada.png", width=Inches(6.2))

if os.path.exists("manual_assets/fig05_ibge_retorno.png"):
    doc.add_picture("manual_assets/fig05_ibge_retorno.png", width=Inches(6.2))

add_callout_box(
    doc,
    "ALGORITMO DE BUSCA INTELIGENTE (FUZZY MATCHING)",
    "Se o nome da cidade tiver pequeno erro de digitação ou acentuação (ex: 'Sao Paulo' em vez de 'São Paulo'), o algoritmo de busca aproximada (Fuzzy Match com threshold de 80%) identifica o município correto e atribui o código IBGE correto sem interromper a execução.",
    "E0F2FE", border_title="💡 RECURSO DE INTELIGÊNCIA"
)

# -----------------------------------------------------------------------------
# CAPÍTULO 4: CRIAR REGIÕES LINCROS
# -----------------------------------------------------------------------------
doc.add_heading("4. Módulo: Criar Regiões Lincros", level=1)
doc.add_paragraph(
    "Este módulo converte a planilha de trabalho do analista no arquivo de importação 'Modelo Região (.xlsx)', estruturando as abas internas 'regioes' e 'localizacoes_atendidas'."
)

doc.add_heading("4.1 Suporte a Regiões por CEP Inicial/Final ou por Código IBGE", level=2)
doc.add_paragraph(
    "O analista pode definir a região usando intervalos de CEP ou códigos IBGE de cidades. O sistema identifica automaticamente as colunas presentes na planilha base e preenche os campos correspondentes na aba 'localizacoes_atendidas':"
)

if os.path.exists("manual_assets/fig06_regiao_retorno.png"):
    doc.add_picture("manual_assets/fig06_regiao_retorno.png", width=Inches(6.5))

add_callout_box(
    doc,
    "ESTRUTURA DAS ABAS GERADAS NO MODELO REGIÃO",
    "• Aba 'regioes': Agrupa os nomes de regiões únicas com o CNPJ da transportadora e insere Ativo = 'VERDADEIRO'.\n"
    "• Aba 'localizacoes_atendidas':\n"
    "   - Se a origem for CEP: Preenche Coluna C (CEP Inicial) e Coluna D (CEP Final).\n"
    "   - Se a origem for Município: Preenche Coluna E (Código IBGE da Cidade).",
    "FEF08A"
)

# -----------------------------------------------------------------------------
# CAPÍTULO 5: GERAR ROTAS LINCROS
# -----------------------------------------------------------------------------
doc.add_heading("5. Módulo: Gerar Rotas Lincros", level=1)
doc.add_paragraph(
    "A ferramenta de rotas cruza o arquivo de regiões preenchido no passo anterior com o ponto de origem da carga (Cidade IBGE ou Nome da Região) para montar a matriz completa de rotas da transportadora."
)

if os.path.exists("manual_assets/fig07_rotas_retorno.png"):
    doc.add_picture("manual_assets/fig07_rotas_retorno.png", width=Inches(6.5))

# -----------------------------------------------------------------------------
# CAPÍTULO 6: FERRAMENTAS AUXILIARES E RESTRIÇÕES POR PESSOAS
# -----------------------------------------------------------------------------
doc.add_heading("6. Ferramentas Auxiliares e Restrições por Pessoas", level=1)

doc.add_heading("6.1 Conversores S/N e STQQS", level=2)
doc.add_paragraph(
    "• Converter S/N: Transforma marcações simples nas colunas dos dias da semana ('S', 'N', '1', '0', 'X') nos valores padrão Lincros ('VERDADEIRO' e 'FALSO').\n"
    "• Converter STQQS: Lê strings codificadas de frequência semanal (ex: 'STQQS..') e desmembra os atendimentos em colunas individuais por dia."
)

doc.add_heading("6.2 Gerador de Restrições por Pessoas (TDE/TAE)", level=2)
doc.add_paragraph(
    "Esta ferramenta lê listas de cadastros de clientes restritos (CNPJ/CPF, Razão Social e Valor) e gera um arquivo compactado `.ZIP` contendo múltiplos arquivos Excel formatados no modelo de Restrições por Pessoas do Lincros, divididos conforme o limite máximo de linhas por arquivo estipulado pelo analista."
)

# Salvar DOCX local e copiar para o Desktop
doc.save(docx_local)
shutil.copy(docx_local, docx_desktop)
print(f"Manual DOCX ultra detalhado criado em: {docx_desktop}")

# =============================================================================
# 2. GERAR DOCUMENTO PDF (REPORTLAB) COMPLETO E DETALHADO
# =============================================================================
styles = getSampleStyleSheet()

style_title = ParagraphStyle('MainTitle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=20, textColor=colors.HexColor('#1a365d'), spaceAfter=6, alignment=1)
style_sub = ParagraphStyle('MainSub', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=11, textColor=colors.HexColor('#4a5568'), spaceAfter=12, alignment=1)
style_h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=13, textColor=colors.HexColor('#1a365d'), spaceBefore=12, spaceAfter=6)
style_h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=10.5, textColor=colors.HexColor('#2b6cb0'), spaceBefore=8, spaceAfter=4)
style_body = ParagraphStyle('Body', parent=styles['Normal'], fontName='Helvetica', fontSize=9, textColor=colors.HexColor('#2d3748'), spaceAfter=5, leading=13)
style_callout_title = ParagraphStyle('CalloutTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=9.5, textColor=colors.HexColor('#854d0e'), spaceAfter=2)
style_callout_body = ParagraphStyle('CalloutBody', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, textColor=colors.HexColor('#422006'), leading=12)

story = []

story.append(Paragraph("MANUAL DO ANALISTA DE LOGÍSTICA", style_title))
story.append(Paragraph("Guia Detalhado de Padronização de Planilhas e Operação no Modelo Lincros", style_sub))
story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#cbd5e0'), spaceAfter=8))

story.append(Paragraph("1. Visão Geral e Fluxo Recomendado de Trabalho", style_h1))
story.append(Paragraph("Para cadastrar uma nova operação ou tabela na Lincros, siga o fluxo padronizado abaixo:", style_body))

for title, desc in passos:
    txt = f"<b>• {title}:</b> {desc}"
    story.append(Paragraph(txt, style_body))

story.append(Spacer(1, 6))

callout_pdf_data = [
    [Paragraph("📌 CONFIGURAÇÕES OBRIGATÓRIAS DA BARRA LATERAL", style_callout_title)],
    [Paragraph("Sempre configure o CNPJ e o Nome da Transportadora Padrão na barra lateral esquerda antes de processar as Regiões ou Rotas para evitar erros de importação na plataforma Lincros!", style_callout_body)]
]
callout_pdf_table = Table(callout_pdf_data, colWidths=[510])
callout_pdf_table.setStyle(TableStyle([
    ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#fef08a')),
    ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#eab308')),
    ('PADDING', (0,0), (-1,-1), 6),
]))
story.append(callout_pdf_table)

story.append(Spacer(1, 8))
story.append(Paragraph("1.1 Estrutura da Planilha Modelo Base Preenchida pelo Analista", style_h2))

if os.path.exists("manual_assets/fig01_base_analista.png"):
    story.append(RLImage("manual_assets/fig01_base_analista.png", width=510, height=135))

story.append(Spacer(1, 8))
story.append(Paragraph("2. Módulo: Cadastro CEP (Lincros)", style_h1))
story.append(Paragraph("Consulta CEPs em tempo real nas APIs dos Correios e gera a planilha oficial 'Modelo CEP Preenchido.xlsx'.", style_body))

if os.path.exists("manual_assets/fig02_cep_entrada.png"):
    story.append(Paragraph("<b>Dados de Entrada (Analista):</b>", style_body))
    story.append(RLImage("manual_assets/fig02_cep_entrada.png", width=490, height=110))

if os.path.exists("manual_assets/fig03_cep_retorno.png"):
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Retorno Gerado pelo Sistema Lincros:</b>", style_body))
    story.append(RLImage("manual_assets/fig03_cep_retorno.png", width=510, height=120))

story.append(Spacer(1, 8))
story.append(Paragraph("3. Módulo: Preencher Códigos IBGE", style_h1))
story.append(Paragraph("Insere os códigos IBGE oficiais de 7 dígitos na planilha base através de busca exata e inteligência aproximada (Fuzzy Match).", style_body))

if os.path.exists("manual_assets/fig05_ibge_retorno.png"):
    story.append(RLImage("manual_assets/fig05_ibge_retorno.png", width=490, height=110))

story.append(Spacer(1, 8))
story.append(Paragraph("4. Módulo: Criar Regiões Lincros", style_h1))
story.append(Paragraph("Estrutura o arquivo oficial Lincros com as abas regioes e localizacoes_atendidas por faixa de CEP ou Código IBGE.", style_body))

if os.path.exists("manual_assets/fig06_regiao_retorno.png"):
    story.append(RLImage("manual_assets/fig06_regiao_retorno.png", width=510, height=115))

story.append(Spacer(1, 8))
story.append(Paragraph("5. Módulo: Gerar Rotas Lincros", style_h1))
story.append(Paragraph("Cruza as regiões atendidas com a origem para construir a matriz de rotas da transportadora.", style_body))

if os.path.exists("manual_assets/fig07_rotas_retorno.png"):
    story.append(RLImage("manual_assets/fig07_rotas_retorno.png", width=510, height=105))

# Build PDF
doc_pdf = SimpleDocTemplate(pdf_local, pagesize=A4, leftMargin=30, rightMargin=30, topMargin=30, bottomMargin=30)
doc_pdf.build(story)

shutil.copy(pdf_local, pdf_desktop)
print(f"Manual PDF ultra detalhado criado em: {pdf_desktop}")
