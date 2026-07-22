import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

import os
import shutil

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak, HRFlowable, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

desktop_path = r"C:\Users\Otávio\Desktop"
os.makedirs(desktop_path, exist_ok=True)

docx_desktop = os.path.join(desktop_path, "Manual_do_Analista_Lincros.docx")
pdf_desktop = os.path.join(desktop_path, "Manual_do_Analista_Lincros.pdf")

docx_local = "Manual_do_Analista_Lincros.docx"
pdf_local = "Manual_do_Analista_Lincros.pdf"

# =============================================================================
# 1. GERAR DOCUMENTO DOCX (WORD)
# =============================================================================
doc = docx.Document()

# Definir margens
sections = doc.sections
for s in sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)

# Funções auxiliares de formatação Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_callout(doc, title, text, bg_color="FEF08A", border_color="EAB308"): # Amarelo marca texto
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    r_title = p.add_run(f"📌 {title}\n")
    r_title.bold = True
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = RGBColor(133, 77, 14)
    
    r_text = p.add_run(text)
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(66, 32, 6)
    
    doc.add_paragraph() # espaçamento

# Cabeçalho Principal
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_t = p_title.add_run("MANUAL DO ANALISTA DE LOGÍSTICA")
r_t.bold = True
r_t.font.size = Pt(22)
r_t.font.color.rgb = RGBColor(26, 54, 93) # Navy blue

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_s = p_sub.add_run("Guia Oficial de Operação e Padronização no Modelo Lincros\n")
r_s.italic = True
r_s.font.size = Pt(13)
r_s.font.color.rgb = RGBColor(74, 85, 104)

doc.add_paragraph()

# Seção 1: Visão Geral
doc.add_heading("1. Fluxo Geral de Trabalho (Passo a Passo)", level=1)
doc.add_paragraph(
    "Para cadastrar uma nova operação de transporte ou tabela de frete/prazos na Lincros, siga o fluxo padronizado abaixo:"
)

items_fluxo = [
    ("Passo 1: Obter a Planilha Modelo Base", "Baixe a planilha limpa clicando em 'Baixar Modelo Base (Vazio)' na barra lateral."),
    ("Passo 2: Identificar Locais (CEP ou IBGE)", "Se você tiver faixas de CEP, use a aba 'Cadastro CEP'. Se tiver nomes de cidades, use a aba 'Preencher IBGE'."),
    ("Passo 3: Montar Prazos e Frequência", "Informe os prazos em dias e marque a frequência semanal ('VERDADEIRO'/'FALSO')."),
    ("Passo 4: Criar Regiões Lincros", "Suba a planilha base na aba 'Criar Região' para gerar a estrutura oficial Lincros ('Regioes.xlsx')."),
    ("Passo 5: Gerar Matriz de Rotas", "Na aba 'Gerar Rotas', cruze o arquivo de regiões com a origem para construir as rotas.")
]

for title, desc in items_fluxo:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(f"{title}: ")
    r1.bold = True
    p.add_run(desc)

add_callout(doc, "REGRA DE OURO PARA ANALISTAS", "Sempre configure o CNPJ e o Nome da Transportadora Padrão na barra lateral esquerda antes de gerar as Regiões ou Rotas para evitar erros de importação!")

# Seção 2: Módulo Cadastro CEP
doc.add_heading("2. Módulo: Cadastro CEP", level=1)
doc.add_paragraph(
    "Esta ferramenta analisa faixas de CEP coladas ou enviadas por planilha e consulta automaticamente o nome oficial no formato Cidade - UF nas APIs públicas dos Correios (ViaCEP, AwesomeAPI, ApiCEP e BrasilAPI)."
)

if os.path.exists("manual_assets/cep_in.png"):
    doc.add_paragraph("Como o Analista preenche os dados de entrada:")
    doc.add_picture("manual_assets/cep_in.png", width=Inches(6.0))

if os.path.exists("manual_assets/cep_out.png"):
    doc.add_paragraph("Como o Sistema Lincros retorna a planilha final:")
    doc.add_picture("manual_assets/cep_out.png", width=Inches(6.0))

add_callout(doc, "MARCA-TEXTO DE VERIFICAÇÃO", "Observe que a coluna Nome é gerada no formato exato 'Cidade - UF', o ID de Localização é mantido vazio e a coluna Ativo é preenchida como 'VERDADEIRO'.")

# Seção 3: Módulo IBGE
doc.add_heading("3. Módulo: Preencher Códigos IBGE", level=1)
doc.add_paragraph(
    "Insere o Código IBGE oficial de 7 dígitos da cidade. Utiliza algoritmos inteligentes de comparação exata e aproximada (Fuzzy Match)."
)

if os.path.exists("manual_assets/ibge_in.png"):
    doc.add_picture("manual_assets/ibge_in.png", width=Inches(6.0))

if os.path.exists("manual_assets/ibge_out.png"):
    doc.add_picture("manual_assets/ibge_out.png", width=Inches(6.0))

# Seção 4: Criar Regiões
doc.add_heading("4. Módulo: Criar Regiões Lincros", level=1)
doc.add_paragraph(
    "Estrutura o arquivo oficial Lincros contendo as abas 'regioes' e 'localizacoes_atendidas'. Suporta criação tanto por Código IBGE quanto por intervalo de CEP (CEP Inicial e CEP Final)."
)

if os.path.exists("manual_assets/reg_in.png"):
    doc.add_picture("manual_assets/reg_in.png", width=Inches(6.0))

if os.path.exists("manual_assets/reg_out.png"):
    doc.add_picture("manual_assets/reg_out.png", width=Inches(6.0))

# Salvar DOCX local e na Área de Trabalho
doc.save(docx_local)
shutil.copy(docx_local, docx_desktop)
print(f"Documento DOCX criado em: {docx_desktop}")


# =============================================================================
# 2. GERAR DOCUMENTO PDF (REPORTLAB)
# =============================================================================
styles = getSampleStyleSheet()

# Personalizar Estilos ReportLab
style_title = ParagraphStyle('MainTitle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=22, textColor=colors.HexColor('#1a365d'), spaceAfter=6, alignment=1)
style_sub = ParagraphStyle('MainSub', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=12, textColor=colors.HexColor('#4a5568'), spaceAfter=18, alignment=1)
style_h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=14, textColor=colors.HexColor('#1a365d'), spaceBefore=14, spaceAfter=8)
style_body = ParagraphStyle('Body', parent=styles['Normal'], fontName='Helvetica', fontSize=10, textColor=colors.HexColor('#2d3748'), spaceAfter=6, leading=14)
style_callout_title = ParagraphStyle('CalloutTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11, textColor=colors.HexColor('#854d0e'), spaceAfter=4)
style_callout_body = ParagraphStyle('CalloutBody', parent=styles['Normal'], fontName='Helvetica', fontSize=9.5, textColor=colors.HexColor('#422006'), leading=13)

story = []

story.append(Paragraph("MANUAL DO ANALISTA DE LOGÍSTICA", style_title))
story.append(Paragraph("Guia Oficial de Operação e Padronização no Modelo Lincros", style_sub))
story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#cbd5e0'), spaceAfter=12))

story.append(Paragraph("1. Fluxo Geral de Trabalho (Passo a Passo)", style_h1))
story.append(Paragraph("Para cadastrar uma nova operação ou tabela na Lincros, siga este fluxo:", style_body))

for title, desc in items_fluxo:
    txt = f"<b>• {title}:</b> {desc}"
    story.append(Paragraph(txt, style_body))

story.append(Spacer(1, 10))

# Callout Amarelo em ReportLab
callout_data = [
    [Paragraph("📌 REGRA DE OURO PARA ANALISTAS", style_callout_title)],
    [Paragraph("Sempre configure o CNPJ e o Nome da Transportadora Padrão na barra lateral esquerda antes de gerar as Regiões ou Rotas para evitar erros de importação!", style_callout_body)]
]
callout_table = Table(callout_data, colWidths=[500])
callout_table.setStyle(TableStyle([
    ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#fef08a')),
    ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#eab308')),
    ('PADDING', (0,0), (-1,-1), 8),
]))
story.append(callout_table)
story.append(Spacer(1, 14))

# Módulo Cadastro CEP
story.append(Paragraph("2. Módulo: Cadastro CEP", style_h1))
story.append(Paragraph("Esta ferramenta analisa faixas de CEP coladas ou enviadas por planilha e consulta automaticamente o nome oficial no formato Cidade - UF nas APIs públicas dos Correios.", style_body))

if os.path.exists("manual_assets/cep_in.png"):
    story.append(Spacer(1, 6))
    story.append(RLImage("manual_assets/cep_in.png", width=480, height=130))

if os.path.exists("manual_assets/cep_out.png"):
    story.append(Spacer(1, 6))
    story.append(RLImage("manual_assets/cep_out.png", width=480, height=130))

story.append(Spacer(1, 10))
story.append(Paragraph("3. Módulo: Preencher Códigos IBGE", style_h1))
story.append(Paragraph("Insere o Código IBGE oficial de 7 dígitos da cidade através de busca exata e aproximada (Fuzzy Match).", style_body))

if os.path.exists("manual_assets/ibge_out.png"):
    story.append(Spacer(1, 6))
    story.append(RLImage("manual_assets/ibge_out.png", width=480, height=130))

story.append(Spacer(1, 10))
story.append(Paragraph("4. Módulo: Criar Regiões Lincros", style_h1))
story.append(Paragraph("Estrutura o arquivo oficial Lincros contendo as abas regioes e localizacoes_atendidas. Suporta criação tanto por Código IBGE quanto por CEP Inicial/Final.", style_body))

if os.path.exists("manual_assets/reg_out.png"):
    story.append(Spacer(1, 6))
    story.append(RLImage("manual_assets/reg_out.png", width=480, height=130))

# Gerar PDF
doc_pdf = SimpleDocTemplate(pdf_local, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
doc_pdf.build(story)

shutil.copy(pdf_local, pdf_desktop)
print(f"Documento PDF criado em: {pdf_desktop}")
